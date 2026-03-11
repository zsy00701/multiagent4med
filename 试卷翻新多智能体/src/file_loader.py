"""
医学试卷翻新多智能体系统 - 文件加载器

支持多种文件格式的读取：.docx, .pdf, .txt
特别注意：医学试卷的选项经常在表格中，必须同时遍历段落和表格。
"""

import logging
from pathlib import Path
from typing import List, Dict, Optional, Generator
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """文档块数据结构"""
    content: str
    source: str
    page: Optional[int] = None
    chunk_id: Optional[int] = None
    metadata: Optional[Dict] = None


class FileLoader:
    """统一文件加载器"""
    
    SUPPORTED_EXTENSIONS = {".docx", ".doc", ".pdf", ".txt"}
    
    def __init__(self):
        self._loaders = {
            ".docx": self._load_docx,
            ".doc": self._load_docx,  # python-docx 也支持 .doc
            ".pdf": self._load_pdf,
            ".txt": self._load_txt,
        }
    
    def load(self, file_path: Path | str) -> str:
        """
        加载文件并返回全部文本内容
        
        Args:
            file_path: 文件路径
        
        Returns:
            文件的全部文本内容
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")
        
        loader = self._loaders.get(ext)
        if loader is None:
            raise ValueError(f"未找到对应的加载器: {ext}")
        
        try:
            content = loader(file_path)
            logger.info(f"成功加载文件: {file_path.name} ({len(content)} 字符)")
            return content
        except Exception as e:
            logger.error(f"加载文件失败 {file_path}: {e}")
            raise
    
    def load_with_chunks(
        self,
        file_path: Path | str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> List[DocumentChunk]:
        """
        加载文件并切分为块
        
        Args:
            file_path: 文件路径
            chunk_size: 每块大小（字符数）
            chunk_overlap: 块之间的重叠（字符数）
        
        Returns:
            文档块列表
        """
        file_path = Path(file_path)
        content = self.load(file_path)
        
        chunks = []
        start = 0
        chunk_id = 0
        content_len = len(content)
        
        while start < content_len:
            end = min(start + chunk_size, content_len)
            chunk_text = content[start:end]
            
            if end < content_len:
                for sep in ("。", ".", "\n", "；", ";"):
                    last_sep = chunk_text.rfind(sep)
                    if last_sep > chunk_size * 0.5:
                        chunk_text = chunk_text[:last_sep + 1]
                        end = start + last_sep + 1
                        break
            
            stripped = chunk_text.strip()
            if stripped:
                chunks.append(DocumentChunk(
                    content=stripped,
                    source=file_path.name,
                    chunk_id=chunk_id,
                    metadata={"file_path": str(file_path), "chunk_size": len(stripped)},
                ))
                chunk_id += 1
            
            next_start = end - chunk_overlap
            if next_start <= start:
                next_start = end
            start = next_start
        
        logger.info(f"文件 {file_path.name} 切分为 {len(chunks)} 个块")
        return chunks
    
    def _load_docx(self, file_path: Path) -> str:
        """
        读取 Word 文档 - 关键：同时遍历段落和表格
        
        医学试卷的选项经常在表格中，必须完整提取！
        """
        from docx import Document
        
        doc = Document(str(file_path))
        content_parts = []
        
        # 建立 element -> 高层对象 的映射（O(n) 替代逐次 O(n) 查找）
        para_map = {para._element: para for para in doc.paragraphs}
        table_map = {table._element: table for table in doc.tables}
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = para_map.get(element)
                if para:
                    text = para.text.strip()
                    if text:
                        content_parts.append(text)
            elif element.tag.endswith('tbl'):
                table = table_map.get(element)
                if table:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        content_parts.append(table_text)
        
        return "\n".join(content_parts)
    
    def _extract_table_text(self, table) -> str:
        """
        提取表格内容为文本
        
        保持行列结构，用于试卷选项的正确提取
        """
        rows_text = []
        
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                # 提取单元格中的所有文本（可能包含多个段落）
                cell_content = []
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        cell_content.append(text)
                
                if cell_content:
                    cells_text.append(" ".join(cell_content))
            
            if cells_text:
                # 用制表符分隔单元格，便于后续解析
                rows_text.append("\t".join(cells_text))
        
        return "\n".join(rows_text)
    
    def _load_pdf(self, file_path: Path) -> str:
        """读取 PDF 文件"""
        from pypdf import PdfReader
        
        reader = PdfReader(str(file_path))
        content_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                # 添加页码标记，便于引用
                content_parts.append(f"[第{page_num}页]\n{text.strip()}")
        
        return "\n\n".join(content_parts)
    
    def _load_txt(self, file_path: Path) -> str:
        """读取纯文本文件"""
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "utf-16"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"无法解码文件 {file_path}，尝试了编码: {encodings}")


def read_docx_full_text(file_path: Path | str) -> str:
    """
    读取 Word 文档的完整文本内容（按文档顺序遍历段落和表格）

    很多医学题的选项在表格里，必须按出现顺序读取。

    Args:
        file_path: Word 文档路径

    Returns:
        文档的完整文本内容
    """
    loader = FileLoader()
    return loader.load(Path(file_path))


def load_knowledge_base(
    directory: Path | str,
    chunk_size: int = 500,
    chunk_overlap: int = 50
) -> Generator[DocumentChunk, None, None]:
    """
    加载知识库目录中的所有文件
    
    Args:
        directory: 知识库目录路径
        chunk_size: 块大小
        chunk_overlap: 重叠大小
    
    Yields:
        文档块
    """
    directory = Path(directory)
    loader = FileLoader()
    
    if not directory.exists():
        logger.warning(f"知识库目录不存在: {directory}")
        return
    
    supported_files = []
    for ext in FileLoader.SUPPORTED_EXTENSIONS:
        supported_files.extend(directory.glob(f"*{ext}"))
    
    logger.info(f"发现 {len(supported_files)} 个知识库文件")
    
    for file_path in sorted(supported_files):
        try:
            chunks = loader.load_with_chunks(
                file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            for chunk in chunks:
                yield chunk
        except Exception as e:
            logger.error(f"加载知识库文件失败 {file_path}: {e}")
            continue
