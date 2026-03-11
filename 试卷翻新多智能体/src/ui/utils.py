"""UI 工具函数"""

import time
from pathlib import Path
from typing import Tuple, List


def validate_file(file) -> Tuple[bool, str]:
    """验证上传的文件

    Args:
        file: Gradio 文件对象

    Returns:
        (是否有效, 错误信息)
    """
    if file is None:
        return False, "未选择文件"

    allowed_extensions = ['.docx', '.doc', '.pdf', '.txt']
    file_path = Path(file.name) if hasattr(file, 'name') else Path(file)
    file_ext = file_path.suffix.lower()

    if file_ext not in allowed_extensions:
        return False, f"不支持的文件格式: {file_ext}，支持的格式: {', '.join(allowed_extensions)}"

    # 检查文件大小 (50MB)
    max_size = 50 * 1024 * 1024
    try:
        if file_path.exists() and file_path.stat().st_size > max_size:
            return False, "文件过大 (最大 50MB)"
    except:
        pass

    return True, "验证通过"


def cleanup_temp_files(temp_dir: Path, max_age_hours: int = 24) -> int:
    """清理过期的临时文件

    Args:
        temp_dir: 临时文件目录
        max_age_hours: 文件最大保留时间(小时)

    Returns:
        清理的文件数量
    """
    if not temp_dir.exists():
        return 0

    cutoff_time = time.time() - (max_age_hours * 3600)
    cleaned_count = 0

    for file in temp_dir.glob("*"):
        if file.is_file():
            try:
                if file.stat().st_mtime < cutoff_time:
                    file.unlink()
                    cleaned_count += 1
            except Exception:
                pass

    return cleaned_count


def format_file_size(size_bytes: int) -> str:
    """格式化文件大小

    Args:
        size_bytes: 字节数

    Returns:
        格式化的字符串 (如 "1.5 MB")
    """
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} TB"


def generate_markdown_without_answers(exam_data: dict) -> str:
    """生成不含答案的 Markdown 试卷

    Args:
        exam_data: 试卷数据字典

    Returns:
        Markdown 文本
    """
    lines = [
        f"# {exam_data['title']}",
        "",
        f"- 版本号: {exam_data.get('version', 'N/A')}",
        f"- 题目总数: {exam_data.get('total_questions', 0)}",
        "",
        "---",
        ""
    ]

    for i, q in enumerate(exam_data.get('questions', []), 1):
        lines.append(f"## 第 {i} 题 ({q.get('type', '未知')})")
        lines.append("")
        lines.append(f"**【考点】** {q.get('knowledge_point', '未知')}")
        lines.append("")
        strategy_text = q.get('renovation_strategy_detail') or q.get('renovation_method')
        if strategy_text:
            lines.append(f"**【翻新策略】** {strategy_text}")
            lines.append("")
        lines.append("**【题干】**")
        lines.append("")
        lines.append(q.get('content', ''))
        lines.append("")

        if q.get('options', {}):
            lines.append("**【选项】**")
        for key, value in q.get('options', {}).items():
            lines.append(f"- {key}. {value}")

        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def generate_markdown_with_answers(exam_data: dict) -> str:
    """生成含答案和解析的 Markdown 试卷

    Args:
        exam_data: 试卷数据字典

    Returns:
        Markdown 文本
    """
    lines = [
        f"# {exam_data['title']}",
        "",
        f"- 版本号: {exam_data.get('version', 'N/A')}",
        f"- 题目总数: {exam_data.get('total_questions', 0)}",
        f"- 通过审核: {exam_data.get('passed_audit_count', 0)}",
        "",
        "---",
        ""
    ]

    for i, q in enumerate(exam_data.get('questions', []), 1):
        lines.append(f"## 第 {i} 题 ({q.get('type', '未知')})")
        lines.append("")
        lines.append(f"**【考点】** {q.get('knowledge_point', '未知')}")
        lines.append("")
        strategy_text = q.get('renovation_strategy_detail') or q.get('renovation_method')
        if strategy_text:
            lines.append(f"**【翻新策略】** {strategy_text}")
            lines.append("")
        lines.append("**【题干】**")
        lines.append("")
        lines.append(q.get('content', ''))
        lines.append("")

        if q.get('options', {}):
            lines.append("**【选项】**")
        for key, value in q.get('options', {}).items():
            lines.append(f"- {key}. {value}")

        lines.append("")
        lines.append(f"**【正确答案】** {q.get('answer', '未知')}")
        lines.append("")

        explanation = q.get('explanation_with_citations') or q.get('explanation', '')
        if explanation:
            lines.append("**【深度解析】**")
            lines.append("")
            lines.append(explanation)
            lines.append("")

        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def markdown_to_pdf(markdown_content: str, output_path: Path) -> bool:
    """将 Markdown 转换为 PDF

    Args:
        markdown_content: Markdown 内容
        output_path: 输出 PDF 路径

    Returns:
        是否成功
    """
    try:
        # 尝试使用 markdown2 + pdfkit
        try:
            import markdown2
            import pdfkit

            html_content = markdown2.markdown(markdown_content)
            html_with_style = f"""
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: "Microsoft YaHei", "SimSun", sans-serif; margin: 40px; }}
                    h1 {{ color: #2c3e50; }}
                    h2 {{ color: #34495e; margin-top: 30px; }}
                    p {{ line-height: 1.6; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            pdfkit.from_string(html_with_style, str(output_path))
            return True
        except ImportError:
            pass

        # 备选方案：使用 reportlab
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # 注册中文字体（如果可用）
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))
                font_name = 'SimSun'
            except:
                font_name = 'Helvetica'

            c = canvas.Canvas(str(output_path), pagesize=A4)
            width, height = A4

            # 简单的文本渲染
            y = height - 50
            for line in markdown_content.split('\n'):
                if y < 50:
                    c.showPage()
                    y = height - 50

                c.setFont(font_name, 12)
                c.drawString(50, y, line[:80])  # 限制每行长度
                y -= 20

            c.save()
            return True
        except ImportError:
            pass

        return False

    except Exception as e:
        print(f"PDF 生成失败: {e}")
        return False


def generate_word_document(exam_data: dict, output_path: Path, include_answers: bool = True) -> bool:
    """生成 Word 文档

    Args:
        exam_data: 试卷数据字典
        output_path: 输出路径
        include_answers: 是否包含答案

    Returns:
        是否成功
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document()

        # 标题
        title = doc.add_heading(exam_data['title'], 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 基本信息
        doc.add_paragraph(f"版本号: {exam_data.get('version', 'N/A')}")
        doc.add_paragraph(f"题目总数: {exam_data.get('total_questions', 0)}")
        if include_answers:
            doc.add_paragraph(f"通过审核: {exam_data.get('passed_audit_count', 0)}")
        doc.add_paragraph()

        # 题目
        for i, q in enumerate(exam_data.get('questions', []), 1):
            # 题号和类型
            heading = doc.add_heading(f"第 {i} 题 ({q.get('type', '未知')})", level=2)

            # 考点
            p = doc.add_paragraph()
            p.add_run('【考点】：').bold = True
            p.add_run(q.get('knowledge_point', '未知'))

            strategy_text = q.get('renovation_strategy_detail') or q.get('renovation_method')
            if strategy_text:
                p = doc.add_paragraph()
                p.add_run('【翻新策略】：').bold = True
                p.add_run(strategy_text)

            # 题干
            p = doc.add_paragraph()
            p.add_run('【题干】：').bold = True
            p.add_run(q.get('content', ''))

            # 选项
            if q.get('options', {}):
                p = doc.add_paragraph()
                p.add_run('【选项】：').bold = True
            for key, value in q.get('options', {}).items():
                doc.add_paragraph(f"{key}. {value}", style='List Bullet')

            # 答案和解析
            if include_answers:
                p = doc.add_paragraph()
                p.add_run('【正确答案】：').bold = True
                run = p.add_run(q.get('answer', '未知'))
                run.font.color.rgb = RGBColor(0, 128, 0)

                explanation = q.get('explanation_with_citations') or q.get('explanation', '')
                if explanation:
                    p = doc.add_paragraph()
                    p.add_run('【深度解析】：').bold = True
                    p.add_run(explanation)

            doc.add_paragraph()  # 空行

        doc.save(str(output_path))
        return True

    except ImportError:
        print("需要安装 python-docx: pip install python-docx")
        return False
    except Exception as e:
        print(f"Word 文档生成失败: {e}")
        return False


def generate_excel_workbook(exam_data: dict, output_path: Path, include_answers: bool = True) -> bool:
    """生成 Excel 题库

    Args:
        exam_data: 试卷数据字典
        output_path: 输出路径
        include_answers: 是否包含答案

    Returns:
        是否成功
    """
    try:
        import pandas as pd

        # 准备数据
        rows = []
        for i, q in enumerate(exam_data.get('questions', []), 1):
            row = {
                '题号': i,
                '题型': q.get('type', '未知'),
                '考点': q.get('knowledge_point', '未知'),
                '题干': q.get('content', ''),
            }

            # 选项
            options = q.get('options', {})
            for key in ['A', 'B', 'C', 'D', 'E', 'F']:
                row[f'选项{key}'] = options.get(key, '')

            if include_answers:
                row['答案'] = q.get('answer', '未知')
                row['解析'] = q.get('explanation_with_citations') or q.get('explanation', '')
                row['通过审核'] = '是' if q.get('passed_audit', False) else '否'
                row['相似度'] = f"{q.get('similarity_score', 0):.2f}"

            rows.append(row)

        # 创建 DataFrame
        df = pd.DataFrame(rows)

        # 保存到 Excel
        with pd.ExcelWriter(str(output_path), engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='题库', index=False)

            # 调整列宽
            worksheet = writer.sheets['题库']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width

        return True

    except ImportError:
        print("需要安装 pandas 和 openpyxl: pip install pandas openpyxl")
        return False
    except Exception as e:
        print(f"Excel 生成失败: {e}")
        return False


def generate_text_file(exam_data: dict, output_path: Path, include_answers: bool = True) -> bool:
    """生成纯文本文件

    Args:
        exam_data: 试卷数据字典
        output_path: 输出路径
        include_answers: 是否包含答案

    Returns:
        是否成功
    """
    try:
        lines = [
            f"{'='*60}",
            exam_data['title'],
            f"{'='*60}",
            "",
            f"版本号: {exam_data.get('version', 'N/A')}",
            f"题目总数: {exam_data.get('total_questions', 0)}",
        ]

        if include_answers:
            lines.append(f"通过审核: {exam_data.get('passed_audit_count', 0)}")

        lines.extend(["", f"{'-'*60}", ""])

        for i, q in enumerate(exam_data.get('questions', []), 1):
            lines.append(f"【第 {i} 题】 {q.get('type', '未知')}")
            lines.append(f"【考点】：{q.get('knowledge_point', '未知')}")
            strategy_text = q.get('renovation_strategy_detail') or q.get('renovation_method')
            if strategy_text:
                lines.append(f"【翻新策略】：{strategy_text}")
            lines.append("")
            lines.append("【题干】：")
            lines.append(q.get('content', ''))
            lines.append("")

            if q.get('options', {}):
                lines.append("【选项】：")
            for key, value in q.get('options', {}).items():
                lines.append(f"  {key}. {value}")

            lines.append("")

            if include_answers:
                lines.append(f"【正确答案】：{q.get('answer', '未知')}")
                explanation = q.get('explanation_with_citations') or q.get('explanation', '')
                if explanation:
                    lines.append("【深度解析】：")
                    lines.append(explanation)
                lines.append("")

            lines.append(f"{'-'*60}")
            lines.append("")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

        return True

    except Exception as e:
        print(f"文本文件生成失败: {e}")
        return False
