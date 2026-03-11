"""
医学试卷翻新多智能体系统 - 格式转换器

将生成的试卷转换为与原始试卷相同的 DOCX 格式。
"""

import logging
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .schemas import GeneratedExamPaper, GeneratedQuestion, QuestionType

logger = logging.getLogger(__name__)


class DocxFormatter:
    """DOCX 格式转换器"""

    def __init__(self):
        self.default_font = "宋体"
        self.title_font_size = 16
        self.heading_font_size = 14
        self.content_font_size = 12

    def convert_to_docx(
        self,
        exam: GeneratedExamPaper,
        output_path: Path,
        include_answers: bool = True,
        include_explanations: bool = False
    ) -> Path:
        """
        将生成的试卷转换为 DOCX 格式

        Args:
            exam: 生成的试卷
            output_path: 输出路径
            include_answers: 是否包含答案
            include_explanations: 是否包含解析

        Returns:
            生成的文件路径
        """
        doc = Document()

        # 设置文档样式
        self._setup_document_style(doc)

        # 添加标题
        self._add_title(doc, exam.original_exam_title or exam.title)
        self._add_header_lines(doc, exam.header_lines)

        # 按原试卷章节与题号添加题目
        current_section = None
        for i, question in enumerate(exam.questions, 1):
            if question.section_title != current_section:
                current_section = question.section_title
                if current_section:
                    self._add_section_heading(doc, current_section)
            self._add_question(doc, question, i)

        # 添加答案部分（如果需要）
        if include_answers:
            self._add_answer_section(doc, exam.questions)

        # 添加解析部分（如果需要）
        if include_explanations:
            self._add_explanation_section(doc, exam.questions)

        # 保存文档
        doc.save(str(output_path))
        logger.info(f"已生成 DOCX 文件: {output_path}")

        return output_path

    def _setup_document_style(self, doc: Document):
        """设置文档样式"""
        # 设置默认字体
        doc.styles['Normal'].font.name = self.default_font
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        doc.styles['Normal'].font.size = Pt(self.content_font_size)

    def _add_title(self, doc: Document, title: str):
        """添加标题"""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_run = title_para.add_run(title)
        title_run.font.name = self.default_font
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        title_run.font.size = Pt(self.title_font_size)
        title_run.font.bold = True

        # 添加空行
        doc.add_paragraph()

    def _add_header_lines(self, doc: Document, header_lines: List[str]):
        """添加原试卷头部信息行，如姓名/工号/成绩。"""
        for line in header_lines:
            doc.add_paragraph(line)
        if header_lines:
            doc.add_paragraph()

    def _add_section_heading(self, doc: Document, section_title: str):
        """添加章节标题，保持原试卷的分区结构。"""
        para = doc.add_paragraph()
        run = para.add_run(section_title)
        run.font.bold = True

    def _add_question(self, doc: Document, question: GeneratedQuestion, number: int):
        """按原试卷结构添加单个题目。"""
        display_number = question.display_number or str(number)
        prefix = f"{display_number}、"

        # 判断题
        if question.type == QuestionType.TRUE_FALSE:
            doc.add_paragraph(f"{prefix}{question.content}                     (    )")
            return

        # 选择题（包括各种类型的选择题）
        if question.type in {QuestionType.A1, QuestionType.A2, QuestionType.A3, QuestionType.A4, QuestionType.SINGLE_CHOICE, QuestionType.MULTIPLE_CHOICE, QuestionType.X, QuestionType.B1}:
            placeholder = "（     ）" if question.type not in {QuestionType.MULTIPLE_CHOICE, QuestionType.X} else "（              ）"
            doc.add_paragraph(f"{prefix}{question.content} {placeholder}")
            for key, value in sorted(question.options.items()):
                option_para = doc.add_paragraph(f"{key}、{value}")
                option_para.paragraph_format.left_indent = Inches(0.2)
            return

        # 简答题和填空题（无选项，保持文字作答形式）
        if question.type in {QuestionType.SHORT_ANSWER, QuestionType.FILL_IN_BLANK}:
            content_lines = [line.strip() for line in question.content.split('\n') if line.strip()]
            if not content_lines:
                doc.add_paragraph(prefix)
                return

            # 第一行带题号
            first_para = doc.add_paragraph(f"{prefix}{content_lines[0]}")
            # 后续行缩进
            for line in content_lines[1:]:
                para = doc.add_paragraph(line)
                para.paragraph_format.left_indent = Inches(0.2)

            # 添加答题空间
            doc.add_paragraph()
            return

        # 其他题型（如案例分析题等）
        content_lines = [line.strip() for line in question.content.split('\n') if line.strip()]
        if not content_lines:
            doc.add_paragraph(prefix)
            return

        first_para = doc.add_paragraph(f"{prefix}{content_lines[0]}")
        for line in content_lines[1:]:
            para = doc.add_paragraph(line)
            para.paragraph_format.left_indent = Inches(0.2)

        doc.add_paragraph()

    def _add_answer_section(self, doc: Document, questions: List[GeneratedQuestion]):
        """添加答案部分"""
        # 分页
        doc.add_page_break()

        # 答案标题
        answer_title = doc.add_paragraph()
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = answer_title.add_run("参考答案")
        title_run.font.name = self.default_font
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        title_run.font.size = Pt(self.heading_font_size)
        title_run.font.bold = True

        doc.add_paragraph()

        # 答案列表
        for i, question in enumerate(questions, 1):
            answer_para = doc.add_paragraph()
            label = f"{question.section_title or '未分区'} {question.display_number or i}"
            answer_para.add_run(f"{label}：").font.bold = True
            answer_para.add_run(question.answer)

    def _add_explanation_section(self, doc: Document, questions: List[GeneratedQuestion]):
        """添加解析部分"""
        # 分页
        doc.add_page_break()

        # 解析标题
        explanation_title = doc.add_paragraph()
        explanation_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = explanation_title.add_run("答案解析")
        title_run.font.name = self.default_font
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        title_run.font.size = Pt(self.heading_font_size)
        title_run.font.bold = True

        doc.add_paragraph()

        # 解析列表
        for i, question in enumerate(questions, 1):
            # 题号
            exp_para = doc.add_paragraph()
            label = f"{question.section_title or '未分区'} {question.display_number or i}"
            exp_para.add_run(f"{label}：").font.bold = True
            exp_para.add_run(f"答案：{question.answer}")

            # 解析内容
            if question.explanation:
                exp_content = doc.add_paragraph(question.explanation)
                exp_content.paragraph_format.left_indent = Inches(0.3)

            doc.add_paragraph()


def convert_exam_to_docx(
    exam: GeneratedExamPaper,
    output_dir: Path,
    include_answers: bool = True,
    include_explanations: bool = False
) -> Path:
    """
    便捷函数：将试卷转换为 DOCX 格式

    Args:
        exam: 生成的试卷
        output_dir: 输出目录
        include_answers: 是否包含答案
        include_explanations: 是否包含解析

    Returns:
        生成的文件路径
    """
    formatter = DocxFormatter()

    # 生成文件名
    filename = f"{exam.title.replace(' ', '_')}.docx"
    output_path = output_dir / filename

    return formatter.convert_to_docx(
        exam,
        output_path,
        include_answers=include_answers,
        include_explanations=include_explanations
    )
